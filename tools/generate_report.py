#!/usr/bin/env python3
"""
HTML Report Generator – Wandelt eine Vertragsanalyse-Markdown-Datei
in ein professionelles, druckfertiges HTML-Dokument um.

Usage:
    python tools/generate_report.py .tmp/vertragsanalyse_*.md
    python tools/generate_report.py .tmp/vertragsanalyse_*.md --output output/
    python tools/generate_report.py .tmp/  (alle .md-Analysen im Ordner)
"""

import sys
import re
from datetime import datetime
from pathlib import Path

sys.path.insert(0, str(Path(__file__).parent))


# ── HTML Template ─────────────────────────────────────────────────────────────

HTML_TEMPLATE = """<!DOCTYPE html>
<html lang="de">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>Vertragsanalyse: {titel}</title>
<style>
  /* ── Reset & Base ── */
  *, *::before, *::after {{ box-sizing: border-box; margin: 0; padding: 0; }}
  html {{ font-size: 15px; }}
  body {{
    font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Helvetica, Arial, sans-serif;
    background: #f0f2f5;
    color: #1a1a2e;
    line-height: 1.65;
  }}

  /* ── Layout ── */
  .page {{ max-width: 900px; margin: 0 auto; padding: 24px 16px 64px; }}

  /* ── Header ── */
  .report-header {{
    background: linear-gradient(135deg, #0f172a 0%, #1e3a5f 100%);
    color: white;
    border-radius: 12px 12px 0 0;
    padding: 36px 40px 32px;
    margin-bottom: 0;
    position: relative;
    overflow: hidden;
  }}
  .report-header::after {{
    content: '';
    position: absolute;
    top: -40px; right: -40px;
    width: 200px; height: 200px;
    border-radius: 50%;
    background: rgba(255,255,255,0.04);
  }}
  .header-top {{
    display: flex;
    justify-content: space-between;
    align-items: flex-start;
    margin-bottom: 20px;
  }}
  .brand {{
    font-size: 11px;
    font-weight: 600;
    letter-spacing: 2px;
    text-transform: uppercase;
    color: rgba(255,255,255,0.5);
    margin-bottom: 8px;
  }}
  .report-title {{
    font-size: 26px;
    font-weight: 700;
    letter-spacing: -0.5px;
    color: white;
    line-height: 1.2;
  }}
  .status-badge {{
    padding: 5px 14px;
    border-radius: 20px;
    font-size: 12px;
    font-weight: 600;
    letter-spacing: 0.5px;
    white-space: nowrap;
    flex-shrink: 0;
    margin-left: 16px;
  }}
  .status-abgeschlossen {{ background: #10b981; color: white; }}
  .status-bearbeitung {{ background: #f59e0b; color: white; }}
  .header-meta {{
    display: flex;
    gap: 24px;
    flex-wrap: wrap;
  }}
  .meta-item {{
    display: flex;
    flex-direction: column;
  }}
  .meta-label {{
    font-size: 10px;
    font-weight: 600;
    text-transform: uppercase;
    letter-spacing: 1px;
    color: rgba(255,255,255,0.4);
    margin-bottom: 2px;
  }}
  .meta-value {{
    font-size: 13px;
    color: rgba(255,255,255,0.85);
    font-weight: 500;
  }}

  /* ── Summary Card ── */
  .summary-card {{
    background: white;
    border: 1px solid #e2e8f0;
    border-top: none;
    padding: 28px 40px;
    display: grid;
    grid-template-columns: 1fr 1fr;
    gap: 20px 48px;
    margin-bottom: 0;
  }}
  .summary-item {{ }}
  .summary-label {{
    font-size: 10px;
    font-weight: 700;
    text-transform: uppercase;
    letter-spacing: 1px;
    color: #94a3b8;
    margin-bottom: 4px;
  }}
  .summary-value {{
    font-size: 14px;
    font-weight: 600;
    color: #1e293b;
  }}
  .summary-value.empty {{ color: #cbd5e1; font-style: italic; font-weight: 400; }}

  /* ── Fristen Box ── */
  .fristen-box {{
    background: #fffbeb;
    border: 1px solid #fde68a;
    border-left: 4px solid #f59e0b;
    border-top: none;
    padding: 20px 40px;
    margin-bottom: 20px;
  }}
  .fristen-header {{
    display: flex;
    align-items: center;
    gap: 8px;
    margin-bottom: 12px;
  }}
  .fristen-icon {{ font-size: 16px; }}
  .fristen-title {{
    font-size: 11px;
    font-weight: 700;
    text-transform: uppercase;
    letter-spacing: 1.5px;
    color: #92400e;
  }}
  .fristen-list {{
    list-style: none;
    display: flex;
    flex-wrap: wrap;
    gap: 8px;
  }}
  .frist-tag {{
    display: inline-flex;
    align-items: center;
    gap: 6px;
    background: white;
    border: 1px solid #fcd34d;
    border-radius: 6px;
    padding: 4px 10px;
    font-size: 12.5px;
    color: #78350f;
    font-weight: 500;
  }}
  .frist-tag .frist-type {{
    font-size: 10px;
    font-weight: 700;
    text-transform: uppercase;
    letter-spacing: 0.5px;
    color: #b45309;
    background: #fef3c7;
    padding: 1px 5px;
    border-radius: 3px;
  }}

  /* ── Section Cards ── */
  .section {{
    background: white;
    border: 1px solid #e2e8f0;
    border-radius: 8px;
    margin-bottom: 12px;
    overflow: hidden;
  }}
  .section-header {{
    padding: 16px 28px;
    border-bottom: 1px solid #f1f5f9;
    display: flex;
    align-items: center;
    gap: 12px;
    background: #fafbfc;
  }}
  .section-number {{
    width: 26px; height: 26px;
    border-radius: 50%;
    background: #1e3a5f;
    color: white;
    font-size: 12px;
    font-weight: 700;
    display: flex;
    align-items: center;
    justify-content: center;
    flex-shrink: 0;
  }}
  .section-title {{
    font-size: 13px;
    font-weight: 700;
    text-transform: uppercase;
    letter-spacing: 1px;
    color: #334155;
  }}
  .section-body {{
    padding: 24px 28px;
  }}

  /* ── Content Styles ── */
  .field-row {{
    display: grid;
    grid-template-columns: 180px 1fr;
    gap: 8px 16px;
    margin-bottom: 10px;
    align-items: baseline;
  }}
  .field-row:last-child {{ margin-bottom: 0; }}
  .field-label {{
    font-size: 12px;
    font-weight: 600;
    color: #64748b;
    padding-top: 1px;
  }}
  .field-value {{
    font-size: 14px;
    color: #1e293b;
  }}
  .field-value.empty {{ color: #cbd5e1; font-style: italic; }}

  .subsection-title {{
    font-size: 12px;
    font-weight: 700;
    text-transform: uppercase;
    letter-spacing: 0.8px;
    color: #64748b;
    margin: 20px 0 10px;
    padding-bottom: 6px;
    border-bottom: 1px solid #f1f5f9;
  }}
  .subsection-title:first-child {{ margin-top: 0; }}

  p, .text-content {{
    font-size: 14px;
    color: #374151;
    line-height: 1.7;
    margin-bottom: 8px;
  }}
  p:last-child {{ margin-bottom: 0; }}

  ul.content-list {{
    list-style: none;
    margin: 0;
    padding: 0;
  }}
  ul.content-list li {{
    font-size: 14px;
    color: #374151;
    padding: 5px 0 5px 20px;
    position: relative;
    border-bottom: 1px solid #f8fafc;
    line-height: 1.6;
  }}
  ul.content-list li:last-child {{ border-bottom: none; }}
  ul.content-list li::before {{
    content: '→';
    position: absolute;
    left: 0;
    color: #94a3b8;
    font-size: 12px;
    top: 6px;
  }}

  /* ── Risiken Section ── */
  .risk-item {{
    background: #fef2f2;
    border: 1px solid #fecaca;
    border-left: 3px solid #ef4444;
    border-radius: 6px;
    padding: 12px 16px;
    margin-bottom: 10px;
    font-size: 14px;
    color: #374151;
    line-height: 1.6;
  }}
  .risk-item:last-child {{ margin-bottom: 0; }}
  .risk-item.medium {{
    background: #fffbeb;
    border-color: #fde68a;
    border-left-color: #f59e0b;
  }}
  .risk-empty {{
    color: #94a3b8;
    font-style: italic;
    font-size: 14px;
  }}

  /* ── Empfehlungen ── */
  .empfehlung-gruppe {{ margin-bottom: 20px; }}
  .empfehlung-gruppe:last-child {{ margin-bottom: 0; }}
  .empfehlung-label {{
    display: inline-flex;
    align-items: center;
    gap: 6px;
    font-size: 11px;
    font-weight: 700;
    text-transform: uppercase;
    letter-spacing: 1px;
    padding: 3px 10px;
    border-radius: 4px;
    margin-bottom: 10px;
  }}
  .label-muss {{ background: #fee2e2; color: #991b1b; }}
  .label-sollte {{ background: #fef3c7; color: #92400e; }}
  .label-kann {{ background: #d1fae5; color: #065f46; }}

  /* ── Rohtext ── */
  .rohtext-box {{
    background: #1e293b;
    border-radius: 6px;
    padding: 20px 24px;
    overflow-x: auto;
  }}
  .rohtext-box pre {{
    font-family: 'SF Mono', 'Fira Code', 'Courier New', monospace;
    font-size: 12px;
    color: #94a3b8;
    white-space: pre-wrap;
    word-break: break-word;
    line-height: 1.7;
  }}

  /* ── Footer ── */
  .report-footer {{
    margin-top: 32px;
    padding: 20px 0 0;
    border-top: 1px solid #e2e8f0;
    display: flex;
    justify-content: space-between;
    align-items: center;
    flex-wrap: gap;
  }}
  .footer-brand {{
    font-size: 12px;
    font-weight: 600;
    color: #94a3b8;
  }}
  .footer-date {{
    font-size: 12px;
    color: #cbd5e1;
  }}

  /* ── Print ── */
  @media print {{
    body {{ background: white; }}
    .page {{ padding: 0; max-width: 100%; }}
    .report-header {{ border-radius: 0; -webkit-print-color-adjust: exact; print-color-adjust: exact; }}
    .section {{ break-inside: avoid; box-shadow: none; }}
    .fristen-box {{ -webkit-print-color-adjust: exact; print-color-adjust: exact; }}
    .risk-item {{ -webkit-print-color-adjust: exact; print-color-adjust: exact; break-inside: avoid; }}
  }}

  @media (max-width: 600px) {{
    .report-header {{ padding: 24px 20px; }}
    .summary-card {{ padding: 20px; grid-template-columns: 1fr; }}
    .fristen-box {{ padding: 16px 20px; }}
    .section-body {{ padding: 16px 20px; }}
    .field-row {{ grid-template-columns: 1fr; gap: 2px; }}
  }}
</style>
</head>
<body>
<div class="page">

  <!-- Header -->
  <div class="report-header">
    <div class="header-top">
      <div>
        <div class="brand">Vertragsanalyse · Automatisiert</div>
        <div class="report-title">{titel}</div>
      </div>
      <span class="status-badge {status_class}">{status}</span>
    </div>
    <div class="header-meta">
      <div class="meta-item">
        <span class="meta-label">Analysedatum</span>
        <span class="meta-value">{analysedatum}</span>
      </div>
      <div class="meta-item">
        <span class="meta-label">Seiten</span>
        <span class="meta-value">{seiten}</span>
      </div>
      <div class="meta-item">
        <span class="meta-label">Analysiert mit</span>
        <span class="meta-value">{analysiert_mit}</span>
      </div>
    </div>
  </div>

  <!-- Summary Card -->
  <div class="summary-card">
    <div class="summary-item">
      <div class="summary-label">Auftraggeber / Partei A</div>
      <div class="summary-value{partei_a_empty}">{partei_a}</div>
    </div>
    <div class="summary-item">
      <div class="summary-label">Vertragstyp</div>
      <div class="summary-value{vertragstyp_empty}">{vertragstyp}</div>
    </div>
    <div class="summary-item">
      <div class="summary-label">Auftragnehmer / Partei B</div>
      <div class="summary-value{partei_b_empty}">{partei_b}</div>
    </div>
    <div class="summary-item">
      <div class="summary-label">Vergütung / Vertragswert</div>
      <div class="summary-value{verguetung_empty}">{verguetung}</div>
    </div>
  </div>

  <!-- Fristen Box -->
  {fristen_box}

  <!-- Sections -->
  {sections_html}

  <!-- Footer -->
  <div class="report-footer">
    <span class="footer-brand">Vertragsmanagement · KI-Analyse</span>
    <span class="footer-date">Erstellt {created_date}</span>
  </div>

</div>
</body>
</html>"""


# ── Markdown Parser ────────────────────────────────────────────────────────────

def parse_md(text: str) -> dict:
    """Parst eine Vertragsanalyse-Markdown-Datei in ein strukturiertes Dict."""

    data = {
        "titel": "",
        "analysedatum": "", "seiten": "", "analysiert_mit": "", "status": "",
        "partei_a": "", "partei_b": "",
        "vertragstyp": "", "leistungsgegenstand": "", "verguetung": "",
        "vertragsbeginn": "", "vertragsende": "", "kuendigungsfrist": "",
        "ausserordentlich": "", "weitere_fristen": "",
        "fristen_auto": [],
        "pflichten_a": "", "pflichten_b": "", "nebenpflichten": "",
        "risiken": "",
        "empf_muss": "", "empf_sollte": "", "empf_kann": "",
        "rohtext": "",
    }

    # Titel aus erster H1
    m = re.search(r"^#\s+Vertragsanalyse:\s*(.+)$", text, re.MULTILINE)
    if m:
        data["titel"] = m.group(1).strip()

    # Header-Felder
    for key, pattern in [
        ("analysedatum", r"\*\*Analysedatum:\*\*\s*(.+)"),
        ("seiten",        r"\*\*Seiten:\*\*\s*(.+)"),
        ("analysiert_mit",r"\*\*Analysiert mit:\*\*\s*(.+)"),
        ("status",        r"\*\*Status:\*\*\s*(.+)"),
    ]:
        m = re.search(pattern, text)
        if m:
            data[key] = m.group(1).strip()

    # Abschnitte extrahieren
    def get_section(nr: int) -> str:
        m = re.search(rf"##\s+{nr}\..+?\n(.*?)(?=\n##\s+\d+\.|\Z)", text, re.DOTALL)
        return m.group(1).strip() if m else ""

    sek1 = get_section(1)
    sek2 = get_section(2)
    sek3 = get_section(3)
    sek4 = get_section(4)
    sek5 = get_section(5)
    sek6 = get_section(6)

    def list_field(section_text: str, label: str) -> str:
        m = re.search(rf"-\s+\*\*{re.escape(label)}[:\*]+\s*(.*)", section_text)
        if m:
            val = m.group(1).strip()
            return "" if val in ("", "-", "–", "n/a", "N/A", "tbd", "TBD") else val
        return ""

    # Abschnitt 1
    data["partei_a"] = list_field(sek1, "Partei A")
    data["partei_b"] = list_field(sek1, "Partei B")

    # Abschnitt 2
    data["vertragstyp"]         = list_field(sek2, "Vertragstyp (BGB-Systematik)")
    data["leistungsgegenstand"] = list_field(sek2, "Leistungsgegenstand")
    data["verguetung"]          = list_field(sek2, "Vergütung / Vertragswert")

    # Abschnitt 3 – Fristen
    data["vertragsbeginn"]   = list_field(sek3, "Vertragsbeginn")
    data["vertragsende"]     = list_field(sek3, "Vertragsende / Laufzeit")
    data["kuendigungsfrist"] = list_field(sek3, "Ordentliche Kündigungsfrist")
    data["ausserordentlich"] = list_field(sek3, "Außerordentliche Kündigung")
    data["weitere_fristen"]  = list_field(sek3, "Weitere Fristen")

    # Auto-erkannte Fristen (aus HTML-Kommentar-Block)
    auto_match = re.search(
        r"<!-- Automatisch erkannt.*?-->(.*?)(?=\n-\s+\*\*Vertragsbeginn|\Z)",
        sek3, re.DOTALL
    )
    if auto_match:
        fristen_text = auto_match.group(1).strip()
        data["fristen_auto"] = [
            line.lstrip("- ").strip()
            for line in fristen_text.splitlines()
            if line.strip().startswith("-")
        ]

    # Abschnitt 4 – Kernpflichten
    def get_subsection(section_text: str, header: str) -> str:
        m = re.search(
            rf"###\s+{re.escape(header)}:?\n(.*?)(?=\n###|\Z)",
            section_text, re.DOTALL
        )
        return m.group(1).strip() if m else ""

    data["pflichten_a"]    = get_subsection(sek4, "Pflichten Partei A")
    data["pflichten_b"]    = get_subsection(sek4, "Pflichten Partei B")
    data["nebenpflichten"] = get_subsection(sek4, "Relevante Nebenpflichten")

    # Abschnitt 5 – Risiken
    data["risiken"] = sek5

    # Abschnitt 6 – Empfehlungen
    def get_empf(section_text: str, label: str) -> str:
        m = re.search(
            rf"\*\*{re.escape(label)}[:\*]+\n(.*?)(?=\n\*\*|\Z)",
            section_text, re.DOTALL
        )
        return m.group(1).strip() if m else ""

    data["empf_muss"]   = get_empf(sek6, "Muss (vor Unterzeichnung)")
    data["empf_sollte"] = get_empf(sek6, "Sollte (empfohlen)")
    data["empf_kann"]   = get_empf(sek6, "Kann (optional)")

    # Rohtext
    m = re.search(r"##\s+7\.\s+Rohtext.*?\n```\n(.*?)```", text, re.DOTALL)
    if m:
        data["rohtext"] = m.group(1).strip()

    return data


# ── HTML Renderer ──────────────────────────────────────────────────────────────

def empty_val(val: str) -> tuple:
    """Gibt (display_val, css_class) zurück."""
    if not val:
        return "Nicht angegeben", " empty"
    return val, ""


def md_to_html_inline(text: str) -> str:
    """Minimal-Markdown: **bold**, Zeilenumbrüche."""
    text = re.sub(r'\*\*(.+?)\*\*', r'<strong>\1</strong>', text)
    text = text.replace('\n', '<br>')
    return text


def render_list_items(text: str) -> str:
    """Wandelt Markdown-Bullet-Liste in HTML <li>-Elemente um."""
    if not text:
        return '<li class="empty-item" style="color:#cbd5e1;font-style:italic;">Nicht angegeben</li>'
    items = []
    for line in text.splitlines():
        line = line.strip()
        if line.startswith("- ") or line.startswith("* "):
            content = md_to_html_inline(line[2:].strip())
            items.append(f"<li>{content}</li>")
        elif line:
            items.append(f"<li>{md_to_html_inline(line)}</li>")
    return "\n".join(items) if items else '<li style="color:#cbd5e1;font-style:italic;">Nicht angegeben</li>'


def render_risiken(text: str) -> str:
    """Rendert Risiken als formatierte Risk-Items."""
    if not text:
        return '<p class="risk-empty">Keine Risiken identifiziert.</p>'

    items = []
    current = []
    for line in text.splitlines():
        stripped = line.strip()
        if stripped.startswith("- ") or stripped.startswith("* "):
            if current:
                items.append(" ".join(current))
            current = [stripped[2:].strip()]
        elif stripped and current:
            current.append(stripped)
        elif stripped and not current:
            current = [stripped]

    if current:
        items.append(" ".join(current))

    if not items:
        return f'<p class="text-content">{md_to_html_inline(text)}</p>'

    html_parts = []
    high_keywords = ["307 BGB", "unwirksam", "nichtig", "Bußgeld", "Haftung", "kritisch"]
    for item in items:
        css_class = "risk-item"
        if not any(kw.lower() in item.lower() for kw in high_keywords):
            css_class += " medium"
        html_parts.append(f'<div class="{css_class}">{md_to_html_inline(item)}</div>')

    return "\n".join(html_parts)


def render_fristen_box(data: dict) -> str:
    """Erstellt die Fristen-Übersichtsbox."""
    items = []

    if data["vertragsbeginn"]:
        items.append(("Datum", data["vertragsbeginn"]))
    if data["vertragsende"]:
        items.append(("Laufzeit", data["vertragsende"]))
    if data["kuendigungsfrist"]:
        items.append(("Kündigung", data["kuendigungsfrist"]))
    if data["ausserordentlich"]:
        items.append(("A.o. Kündigung", data["ausserordentlich"]))
    if data["weitere_fristen"]:
        items.append(("Weitere", data["weitere_fristen"]))

    # Auto-erkannte Fristen ergänzen
    for frist in data["fristen_auto"][:4]:
        m = re.match(r"(.+?)\s+\((.+?)\)", frist)
        if m:
            items.append((m.group(2), m.group(1)))
        else:
            items.append(("Frist", frist))

    if not items:
        return ""

    tags = []
    for typ, wert in items:
        tags.append(
            f'<li class="frist-tag">'
            f'<span class="frist-type">{typ}</span>{wert}'
            f'</li>'
        )

    return f"""
  <div class="fristen-box">
    <div class="fristen-header">
      <span class="fristen-icon">⏱</span>
      <span class="fristen-title">Fristen &amp; Termine auf einen Blick</span>
    </div>
    <ul class="fristen-list">
      {"".join(tags)}
    </ul>
  </div>"""


def render_sections(data: dict) -> str:
    parts = []

    # ── Abschnitt 1: Parteien ──
    pa, pa_c = empty_val(data["partei_a"])
    pb, pb_c = empty_val(data["partei_b"])
    parts.append(f"""
  <div class="section">
    <div class="section-header">
      <span class="section-number">1</span>
      <span class="section-title">Vertragsparteien</span>
    </div>
    <div class="section-body">
      <div class="field-row">
        <span class="field-label">Partei A (Auftraggeber)</span>
        <span class="field-value{pa_c}">{pa}</span>
      </div>
      <div class="field-row">
        <span class="field-label">Partei B (Auftragnehmer)</span>
        <span class="field-value{pb_c}">{pb}</span>
      </div>
    </div>
  </div>""")

    # ── Abschnitt 2: Typ & Gegenstand ──
    vt, vt_c = empty_val(data["vertragstyp"])
    lg, lg_c = empty_val(data["leistungsgegenstand"])
    vg, vg_c = empty_val(data["verguetung"])
    parts.append(f"""
  <div class="section">
    <div class="section-header">
      <span class="section-number">2</span>
      <span class="section-title">Vertragstyp &amp; Gegenstand</span>
    </div>
    <div class="section-body">
      <div class="field-row">
        <span class="field-label">Vertragstyp (BGB)</span>
        <span class="field-value{vt_c}">{vt}</span>
      </div>
      <div class="field-row">
        <span class="field-label">Leistungsgegenstand</span>
        <span class="field-value{lg_c}">{lg}</span>
      </div>
      <div class="field-row">
        <span class="field-label">Vergütung / Wert</span>
        <span class="field-value{vg_c}">{vg}</span>
      </div>
    </div>
  </div>""")

    # ── Abschnitt 3: Laufzeit & Fristen ──
    vb, vb_c = empty_val(data["vertragsbeginn"])
    ve, ve_c = empty_val(data["vertragsende"])
    kf, kf_c = empty_val(data["kuendigungsfrist"])
    ak, ak_c = empty_val(data["ausserordentlich"])
    wf, wf_c = empty_val(data["weitere_fristen"])
    parts.append(f"""
  <div class="section">
    <div class="section-header">
      <span class="section-number">3</span>
      <span class="section-title">Laufzeit &amp; Fristen</span>
    </div>
    <div class="section-body">
      <div class="field-row">
        <span class="field-label">Vertragsbeginn</span>
        <span class="field-value{vb_c}">{vb}</span>
      </div>
      <div class="field-row">
        <span class="field-label">Vertragsende / Laufzeit</span>
        <span class="field-value{ve_c}">{ve}</span>
      </div>
      <div class="field-row">
        <span class="field-label">Ordentliche Kündigung</span>
        <span class="field-value{kf_c}">{kf}</span>
      </div>
      <div class="field-row">
        <span class="field-label">Außerordentliche Kündigung</span>
        <span class="field-value{ak_c}">{ak}</span>
      </div>
      <div class="field-row">
        <span class="field-label">Weitere Fristen</span>
        <span class="field-value{wf_c}">{wf}</span>
      </div>
    </div>
  </div>""")

    # ── Abschnitt 4: Kernpflichten ──
    parts.append(f"""
  <div class="section">
    <div class="section-header">
      <span class="section-number">4</span>
      <span class="section-title">Kernpflichten</span>
    </div>
    <div class="section-body">
      <div class="subsection-title">Pflichten Partei A</div>
      <ul class="content-list">{render_list_items(data["pflichten_a"])}</ul>
      <div class="subsection-title">Pflichten Partei B</div>
      <ul class="content-list">{render_list_items(data["pflichten_b"])}</ul>
      <div class="subsection-title">Relevante Nebenpflichten</div>
      <ul class="content-list">{render_list_items(data["nebenpflichten"])}</ul>
    </div>
  </div>""")

    # ── Abschnitt 5: Risiken ──
    parts.append(f"""
  <div class="section">
    <div class="section-header">
      <span class="section-number">5</span>
      <span class="section-title">Rechtliche Risiken &amp; Auffälligkeiten</span>
    </div>
    <div class="section-body">
      {render_risiken(data["risiken"])}
    </div>
  </div>""")

    # ── Abschnitt 6: Handlungsempfehlungen ──
    muss   = render_list_items(data["empf_muss"])
    sollte = render_list_items(data["empf_sollte"])
    kann   = render_list_items(data["empf_kann"])
    parts.append(f"""
  <div class="section">
    <div class="section-header">
      <span class="section-number">6</span>
      <span class="section-title">Handlungsempfehlungen</span>
    </div>
    <div class="section-body">
      <div class="empfehlung-gruppe">
        <div class="empfehlung-label label-muss">⛔ Muss – vor Unterzeichnung</div>
        <ul class="content-list">{muss}</ul>
      </div>
      <div class="empfehlung-gruppe">
        <div class="empfehlung-label label-sollte">⚠ Sollte – empfohlen</div>
        <ul class="content-list">{sollte}</ul>
      </div>
      <div class="empfehlung-gruppe">
        <div class="empfehlung-label label-kann">✓ Kann – optional</div>
        <ul class="content-list">{kann}</ul>
      </div>
    </div>
  </div>""")

    # ── Rohtext (kollabiert) ──
    if data["rohtext"]:
        import html as html_lib
        rohtext_escaped = html_lib.escape(data["rohtext"])
        parts.append(f"""
  <div class="section">
    <div class="section-header">
      <span class="section-number">7</span>
      <span class="section-title">Rohtext (Quelle)</span>
    </div>
    <div class="section-body">
      <div class="rohtext-box">
        <pre>{rohtext_escaped}</pre>
      </div>
    </div>
  </div>""")

    return "\n".join(parts)


# ── Haupt-Render-Funktion ─────────────────────────────────────────────────────

def render_html(md_text: str) -> str:
    data = parse_md(md_text)

    status = data["status"] or "In Bearbeitung"
    status_class = (
        "status-abgeschlossen" if "abgeschlossen" in status.lower()
        else "status-bearbeitung"
    )

    pa, pa_c = empty_val(data["partei_a"])
    pb, pb_c = empty_val(data["partei_b"])
    vt, vt_c = empty_val(data["vertragstyp"])
    vg, vg_c = empty_val(data["verguetung"])

    return HTML_TEMPLATE.format(
        titel           = data["titel"] or "Vertragsanalyse",
        status          = status,
        status_class    = status_class,
        analysedatum    = data["analysedatum"] or "–",
        seiten          = data["seiten"] or "–",
        analysiert_mit  = data["analysiert_mit"] or "–",
        partei_a        = pa, partei_a_empty = pa_c,
        partei_b        = pb, partei_b_empty = pb_c,
        vertragstyp     = vt, vertragstyp_empty = vt_c,
        verguetung      = vg, verguetung_empty = vg_c,
        fristen_box     = render_fristen_box(data),
        sections_html   = render_sections(data),
        created_date    = datetime.now().strftime("%d.%m.%Y"),
    )


# ── Importierbare Funktion ────────────────────────────────────────────────────

def generiere_report(md_pfad: Path, output_dir: Path = None) -> Path:
    """Generiert einen HTML-Report aus einer .md-Analysedatei. Gibt den Output-Pfad zurück."""
    md_text = md_pfad.read_text(encoding="utf-8")
    html = render_html(md_text)
    if output_dir:
        output_dir.mkdir(parents=True, exist_ok=True)
        out = output_dir / (md_pfad.stem + ".html")
    else:
        out = md_pfad.with_suffix(".html")
    out.write_text(html, encoding="utf-8")
    return out


# ── CLI ───────────────────────────────────────────────────────────────────────

def main():
    args = sys.argv[1:]
    if not args or args[0] in ("-h", "--help"):
        print("Verwendung:")
        print("  python tools/generate_report.py analyse.md")
        print("  python tools/generate_report.py analyse_a.md analyse_b.md")
        print("  python tools/generate_report.py output/       (alle vertragsanalyse_*.md im Ordner)")
        print("  python tools/generate_report.py output/ --output output/reports/")
        sys.exit(0)

    # --output parsen und aus args entfernen
    output_dir = None
    if "--output" in args:
        idx = args.index("--output")
        if idx + 1 < len(args):
            output_dir = Path(args[idx + 1])
        args = args[:idx] + args[idx + 2:]

    # Alle .md-Dateien sammeln (Datei oder Ordner, mehrere Argumente)
    md_dateien = []
    for arg in args:
        p = Path(arg)
        if p.is_dir():
            gefunden = sorted(p.glob("vertragsanalyse_*.md"))
            if not gefunden:
                print(f"Warnung: Keine vertragsanalyse_*.md in {p}")
            md_dateien.extend(gefunden)
        elif p.is_file():
            md_dateien.append(p)
        else:
            print(f"Warnung: Nicht gefunden – {p}")

    if not md_dateien:
        print("Keine .md-Dateien gefunden.")
        sys.exit(1)

    erzeugte = []
    for md_pfad in md_dateien:
        try:
            out = generiere_report(md_pfad, output_dir)
            erzeugte.append(out)
            print(f"✓ {out}")
        except Exception as e:
            print(f"⚠ {md_pfad.name} übersprungen: {e}")

    print(f"\nFertig: {len(erzeugte)} Report(s) erstellt.")


def generiere_docx(md_pfad, output_dir=None):
    """
    Wandelt eine Vertragsanalyse-Markdown-Datei in ein DOCX-Dokument um.
    Gibt den Pfad zur erstellten .docx-Datei zurück.
    """
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        raise ImportError("python-docx fehlt. Lösung: pip install python-docx")

    md_pfad = Path(md_pfad)
    text = md_pfad.read_text(encoding="utf-8")

    if output_dir is None:
        output_dir = md_pfad.parent
    output_dir = Path(output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)

    stamm = re.sub(r"_\d{8}_\d{6}$", "", md_pfad.stem)
    datum_str = datetime.now().strftime("%Y%m%d_%H%M%S")
    docx_pfad = output_dir / f"{stamm}_{datum_str}.docx"

    doc = Document()

    # ── Seitenränder ──
    section = doc.sections[0]
    section.left_margin   = Inches(1.0)
    section.right_margin  = Inches(1.0)
    section.top_margin    = Inches(1.0)
    section.bottom_margin = Inches(1.0)

    # ── Styles ──
    def stil_ueberschrift1(para):
        para.style = doc.styles["Heading 1"]
        run = para.runs[0] if para.runs else para.add_run(para.text)
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)

    def stil_ueberschrift2(para):
        para.style = doc.styles["Heading 2"]
        run = para.runs[0] if para.runs else para.add_run(para.text)
        run.font.size = Pt(13)
        run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)

    # ── Titel-Block ──
    titel_match = re.search(r"#\s+Vertragsanalyse:\s+(.+)", text)
    titel = titel_match.group(1).strip() if titel_match else md_pfad.stem

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Vertragsanalyse")
    run.font.size = Pt(20)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run(titel)
    run2.font.size = Pt(14)
    run2.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)

    # Header-Felder (Analysedatum, Status, Modell)
    for feld in ["Analysedatum", "Status", "Analysiert mit"]:
        m = re.search(rf"\*\*{feld}:\*\*\s*(.+)", text)
        if m:
            p = doc.add_paragraph()
            run = p.add_run(f"{feld}: ")
            run.font.bold = True
            run.font.size = Pt(10)
            run2 = p.add_run(m.group(1).strip())
            run2.font.size = Pt(10)

    doc.add_paragraph()  # Leerzeile

    # ── Abschnitte parsen und ausgeben ──
    abschnitt_pattern = re.compile(
        r"^##\s+(\d+\.\s+.+?)$\n(.*?)(?=^##\s+\d+\.|\Z)",
        re.MULTILINE | re.DOTALL,
    )

    for match in abschnitt_pattern.finditer(text):
        abschnitt_titel = match.group(1).strip()
        abschnitt_inhalt = match.group(2).strip()

        # Abschnitt-Überschrift
        p = doc.add_paragraph()
        stil_ueberschrift2(p)
        p.clear()
        p.add_run(abschnitt_titel).font.size = Pt(13)

        # Inhalt zeilenweise
        for zeile in abschnitt_inhalt.splitlines():
            zeile = zeile.strip()
            if not zeile or zeile == "---":
                continue

            # ### Unterüberschrift
            if zeile.startswith("### "):
                p = doc.add_paragraph()
                run = p.add_run(zeile[4:])
                run.font.bold = True
                run.font.size = Pt(11)
                run.font.color.rgb = RGBColor(0x37, 0x41, 0x51)

            # Listeneintrag
            elif zeile.startswith("- "):
                inhalt = zeile[2:]
                # Fettdruck **..:** am Anfang erkennen
                fett_match = re.match(r"\*\*(.+?):\*\*\s*(.*)", inhalt)
                p = doc.add_paragraph(style="List Bullet")
                if fett_match:
                    run = p.add_run(fett_match.group(1) + ": ")
                    run.font.bold = True
                    run.font.size = Pt(10)
                    rest = re.sub(r"\*+", "", fett_match.group(2))
                    p.add_run(rest).font.size = Pt(10)
                else:
                    clean = re.sub(r"\*+", "", inhalt)
                    p.add_run(clean).font.size = Pt(10)

            # Normaler Text
            else:
                clean = re.sub(r"\*+", "", zeile)
                if clean:
                    p = doc.add_paragraph(clean)
                    p.runs[0].font.size = Pt(10) if p.runs else None

        doc.add_paragraph()  # Abstand nach Abschnitt

    # ── Footer ──
    doc.add_paragraph("─" * 60)
    p = doc.add_paragraph(
        f"Erstellt mit Vertragsmanagement v2.0 · {datetime.now().strftime('%d.%m.%Y')}"
    )
    p.runs[0].font.size = Pt(9)
    p.runs[0].font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)

    doc.save(str(docx_pfad))
    return str(docx_pfad)


if __name__ == "__main__":
    main()
